"""
Export module for ScriptForge
Handles exporting scripts to various formats
"""

import os
import json
from datetime import datetime
from typing import Dict, Any, Optional
import config

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def export_script(script_text: str, metadata: Dict[str, Any], format: str = "txt") -> Optional[str]:
    """
    Export script to specified format
    
    Args:
        script_text: The script text to export
        metadata: Script metadata including parameters
        format: Export format ("txt", "docx", "json")
        
    Returns:
        Path to exported file or None if failed
    """
    
    if format not in config.EXPORT_FORMATS:
        print(f"Unsupported export format: {format}")
        return None
    
    # Create export directory if it doesn't exist
    os.makedirs(config.EXPORT_DIR, exist_ok=True)
    
    # Generate filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    topic_slug = metadata.get("params", {}).get("topic", "script")[:50].replace(" ", "_").lower()
    filename = f"scriptforge_{topic_slug}_{timestamp}"
    
    try:
        if format == "txt":
            return export_to_txt(script_text, metadata, filename)
        elif format == "docx":
            return export_to_docx(script_text, metadata, filename)
        elif format == "json":
            return export_to_json(script_text, metadata, filename)
            
    except Exception as e:
        print(f"Error exporting script to {format}: {e}")
        return None

def export_to_txt(script_text: str, metadata: Dict[str, Any], base_filename: str) -> str:
    """
    Export script to plain text file
    
    Args:
        script_text: The script text
        metadata: Script metadata
        base_filename: Base filename without extension
        
    Returns:
        Path to exported file
    """
    filepath = os.path.join(config.EXPORT_DIR, f"{base_filename}.txt")
    
    with open(filepath, "w", encoding="utf-8") as f:
        # Write header
        f.write("=" * 60 + "\n")
        f.write(f"SCRIPTFORGE - AI Video Script Generator\n")
        f.write(f"Generated: {metadata.get('generated_at', 'Unknown')}\n")
        f.write("=" * 60 + "\n\n")
        
        # Write metadata
        params = metadata.get("params", {})
        f.write("SCRIPT METADATA:\n")
        f.write("-" * 40 + "\n")
        f.write(f"Topic: {params.get('topic', 'N/A')}\n")
        f.write(f"Video Type: {params.get('video_type', 'N/A')}\n")
        f.write(f"Video Format: {params.get('video_format', 'N/A')}\n")
        f.write(f"Target Length: {params.get('target_length', 'N/A')} ")
        f.write("minutes" if params.get('video_type') == 'YouTube' else "seconds")
        f.write("\n")
        f.write(f"Tone: {params.get('tone', 'N/A')}\n")
        f.write(f"Audience: {', '.join(params.get('audience', []))}\n")
        f.write(f"Model Used: {metadata.get('model_used', 'Unknown')}\n")
        f.write("\n" + "=" * 60 + "\n\n")
        
        # Write script
        f.write("SCRIPT CONTENT:\n")
        f.write("-" * 40 + "\n\n")
        f.write(script_text)
        f.write("\n\n" + "=" * 60 + "\n")
        f.write("End of Script\n")
        f.write("=" * 60 + "\n")
    
    return filepath

def export_to_docx(script_text: str, metadata: Dict[str, Any], base_filename: str) -> Optional[str]:
    """
    Export script to Microsoft Word document
    
    Args:
        script_text: The script text
        metadata: Script metadata
        base_filename: Base filename without extension
        
    Returns:
        Path to exported file or None if docx not available
    """
    if not DOCX_AVAILABLE:
        print("python-docx not available. Install with: pip install python-docx")
        return None
    
    filepath = os.path.join(config.EXPORT_DIR, f"{base_filename}.docx")
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('ScriptForge - AI Video Script', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph(f"Generated: {metadata.get('generated_at', 'Unknown')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Add metadata section
    doc.add_heading('Script Metadata', level=1)
    
    params = metadata.get("params", {})
    metadata_table = doc.add_table(rows=7, cols=2)
    metadata_table.style = 'Light Shading Accent 1'
    
    # Fill metadata table
    rows_data = [
        ("Topic:", params.get('topic', 'N/A')),
        ("Video Type:", params.get('video_type', 'N/A')),
        ("Video Format:", params.get('video_format', 'N/A')),
        ("Target Length:", f"{params.get('target_length', 'N/A')} "
         f"{'minutes' if params.get('video_type') == 'YouTube' else 'seconds'}"),
        ("Tone:", params.get('tone', 'N/A')),
        ("Audience:", ', '.join(params.get('audience', []))),
        ("Model Used:", metadata.get('model_used', 'Unknown'))
    ]
    
    for i, (label, value) in enumerate(rows_data):
        metadata_table.cell(i, 0).text = label
        metadata_table.cell(i, 1).text = str(value)
    
    doc.add_page_break()
    
    # Add script content
    doc.add_heading('Script Content', level=1)
    
    # Split script into lines and add with formatting
    lines = script_text.split('\n')
    for line in lines:
        if line.strip():
            # Check if line is a section header (contains brackets)
            if '[' in line and ']' in line:
                # It's a section header
                para = doc.add_paragraph(line)
                para.style = 'Heading 2'
            elif line.startswith('Hook:') or line.startswith('Preview:') or line.startswith('Summary:') or line.startswith('CTA:'):
                # It's a labeled section
                para = doc.add_paragraph(line)
                para.style = 'Intense Quote'
            elif ':' in line and not line.startswith(' '):
                # It's a bullet point or similar
                para = doc.add_paragraph(line)
            else:
                # Regular content
                para = doc.add_paragraph(line)
                para.style = 'Normal'
        else:
            # Empty line
            doc.add_paragraph()
    
    # Add footer
    doc.add_page_break()
    footer = doc.add_paragraph("Generated by ScriptForge - AI Video Script Generator")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    doc.save(filepath)
    
    return filepath

def export_to_json(script_text: str, metadata: Dict[str, Any], base_filename: str) -> str:
    """
    Export script to JSON format
    
    Args:
        script_text: The script text
        metadata: Script metadata
        base_filename: Base filename without extension
        
    Returns:
        Path to exported file
    """
    filepath = os.path.join(config.EXPORT_DIR, f"{base_filename}.json")
    
    export_data = {
        "metadata": {
            "app": config.APP_NAME,
            "version": config.APP_VERSION,
            "generated_at": metadata.get("generated_at"),
            "model_used": metadata.get("model_used"),
            "exported_at": datetime.now().isoformat()
        },
        "parameters": metadata.get("params", {}),
        "script": script_text
    }
    
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(export_data, f, indent=2, ensure_ascii=False)
    
    return filepath